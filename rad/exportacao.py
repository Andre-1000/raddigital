"""
Exportacao do RAD — mensagem "Copiar msg" (RG-EXP-013) e PDF (RG-EXP-
001 a 012), a partir de um RAD ja sincronizado.

Decisao de negocio (17/07/2026): o template original da EFD (secao
3.13) NAO muda de estrutura, mas todos os campos adicionados nas
mudancas de negocio recentes (OS/SA, Responsavel Atividade, Operador
CCM, Descricao Tecnica da Atividade, Equipes Envolvidas) precisam
aparecer nele. _campos_do_relatorio() e o unico lugar que define essa
lista -- tanto a mensagem quanto o PDF sao gerados a partir dela, para
nunca divergirem entre si.

Nota de arquitetura: a EFD descreve a exportacao como uma acao
client-side, offline, executada ANTES da sincronizacao (RG-EXP-002/011)
pelo proprio tecnico que preencheu o RAD. Esta implementacao e um
recurso adicional do backend, operando sobre um RAD JA sincronizado --
util para Supervisor/Administrador na tela de consulta, e para o
proprio criador do RAD reobter o arquivo depois. Nao substitui a
exportacao offline eventualmente implementada no cliente; quando essa
existir, deve reutilizar esta mesma lista de campos como referencia.

Campos desabilitados (app configuracoes) sao omitidos tanto da mensagem
quanto do PDF, seguindo a mesma regra do resto do sistema.
"""
from io import BytesIO

from configuracoes.servicos import campos_desabilitados

NAO_APLICAVEL = 'N/A'


def _ou_na(valor):
    if valor is None or (isinstance(valor, str) and not valor.strip()):
        return NAO_APLICAVEL
    return str(valor)


def _lista_ou_na(valores):
    valores = list(valores)
    return ', '.join(valores) if valores else NAO_APLICAVEL


def _servicos_texto(rad):
    nomes = list(rad.servicos.values_list('servico__nome', flat=True))
    texto = ', '.join(nomes) if nomes else NAO_APLICAVEL
    if rad.outros_servico_desc:
        texto += f' ({rad.outros_servico_desc})'
    return texto


def _motivo_atrasos_texto(rad):
    partes = []
    if rad.atraso_inicio and rad.motivo_atraso_inicio:
        motivo = rad.motivo_atraso_inicio.nome
        if rad.desc_motivo_atraso_inicio:
            motivo += f': {rad.desc_motivo_atraso_inicio}'
        partes.append(f'Início — {motivo}')
    if rad.atraso_termino and rad.motivo_atraso_termino:
        motivo = rad.motivo_atraso_termino.nome
        if rad.desc_motivo_atraso_termino:
            motivo += f': {rad.desc_motivo_atraso_termino}'
        partes.append(f'Término — {motivo}')
    return '; '.join(partes) if partes else NAO_APLICAVEL


def _responsaveis_texto(rad):
    return _lista_ou_na(rad.colaboradores.values_list('nome', flat=True))


def _campos_do_relatorio(rad):
    """
    Lista ordenada (chave_campo, rotulo, valor) — fonte unica de
    verdade tanto para a mensagem quanto para o PDF. chave_campo
    corresponde as chaves de configuracoes.CampoFormulario, usada para
    filtrar campos desabilitados.
    """
    return [
        ('servicos', 'Atividade', _servicos_texto(rad)),
        ('data_preenchimento', 'Data', rad.data_preenchimento.strftime('%d/%m/%Y')),
        ('numero_os', 'OS', str(rad.numero_os)),
        ('numero_sa', 'N° SA', rad.numero_sa),
        ('numero_falha', 'Falha', _ou_na(rad.numero_falha)),
        ('local_inicial', 'Local', f'{rad.local_inicial.sigla}/{rad.local_final.sigla}'),
        ('linhas', 'Linha', _lista_ou_na(rad.linhas.values_list('linha__nome', flat=True))),
        ('vias', 'Via', _lista_ou_na(rad.vias.values_list('via__nome', flat=True))),
        (
            'equipes',
            'Equipes Envolvidas',
            _lista_ou_na(rad.equipes.values_list('equipe__nome', flat=True)),
        ),
        ('km_poste', 'Km/Poste', _ou_na(rad.km_poste)),
        (
            'hora_prog_inicio',
            'Horário programado',
            f'{rad.hora_prog_inicio:%H:%M} a {rad.hora_prog_termino:%H:%M}',
        ),
        ('hora_real_inicio', 'Início', f'{rad.hora_real_inicio:%H:%M}'),
        ('hora_real_termino', 'Término', f'{rad.hora_real_termino:%H:%M}'),
        ('servicos', 'Serviços realizados', _servicos_texto(rad)),  # duplicado, igual a EFD 3.13
        ('materiais_utilizados', 'Equipamentos utilizados', _ou_na(rad.materiais_utilizados)),
        ('motivo_atraso_inicio', 'Motivo dos atrasos', _motivo_atrasos_texto(rad)),
        ('colaboradores', 'Responsável', _responsaveis_texto(rad)),
        ('responsavel_atividade', 'Responsável Atividade', _ou_na(rad.responsavel_atividade)),
        ('operador_ccm', 'Operador CCM', _ou_na(rad.operador_ccm)),
        (
            'descricao_tecnica_atividade',
            'Descrição Técnica da Atividade',
            _ou_na(rad.descricao_tecnica_atividade),
        ),
        ('observacoes_gerais', 'Observação Geral', _ou_na(rad.observacoes_gerais)),
    ]


def gerar_mensagem_copiar(rad):
    """RG-EXP-013: mensagem estruturada para a área de transferência."""
    desabilitados = campos_desabilitados()
    campos = _campos_do_relatorio(rad)
    corpo = '\n\n'.join(
        f'{rotulo}: {valor}' for chave, rotulo, valor in campos if chave not in desabilitados
    )
    return f'RAD - (Relatório de Atividade Diária)\n\n{corpo}'


def gerar_pdf_bytes(rad):
    """RG-EXP-001 a 012: gera o PDF do RAD com todos os campos aplicaveis."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    desabilitados = campos_desabilitados()
    campos = _campos_do_relatorio(rad)
    styles = getSampleStyleSheet()

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )

    story = [
        Paragraph('RAD — Relatório de Atividade Diária', styles['Title']),
        Spacer(1, 6),
        Paragraph(
            f'RAD N° {rad.numero_rad} — Status: {rad.get_status_display().capitalize()}',
            styles['Normal'],
        ),
        Spacer(1, 12),
    ]

    estilo_valor = styles['Normal']
    estilo_valor.wordWrap = 'CJK'
    linhas_tabela = [
        [Paragraph(f'<b>{rotulo}</b>', estilo_valor), Paragraph(valor, estilo_valor)]
        for chave, rotulo, valor in campos
        if chave not in desabilitados
    ]
    tabela = Table(linhas_tabela, colWidths=[5 * cm, 11 * cm])
    tabela.setStyle(
        TableStyle(
            [
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.whitesmoke),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(tabela)

    mostra_fotos_intervencao = 'fotos_intervencao_verificada' not in desabilitados
    mostra_fotos_acao = 'fotos_acao_realizada' not in desabilitados
    mostra_pdf_anexo = 'pdf' not in desabilitados
    if mostra_fotos_intervencao or mostra_fotos_acao or mostra_pdf_anexo:
        story.append(Spacer(1, 16))
        story.append(Paragraph('Anexos', styles['Heading2']))
        if mostra_fotos_intervencao:
            qtd = rad.anexos.filter(
                tipo_arquivo='foto', categoria_foto='intervencao_verificada'
            ).count()
            story.append(Paragraph(f'Fotos — Intervenção Verificada: {qtd}', styles['Normal']))
        if mostra_fotos_acao:
            qtd = rad.anexos.filter(tipo_arquivo='foto', categoria_foto='acao_realizada').count()
            story.append(Paragraph(f'Fotos — Ação Realizada: {qtd}', styles['Normal']))
        if mostra_pdf_anexo:
            tem_pdf = rad.anexos.filter(tipo_arquivo='pdf').exists()
            story.append(Paragraph(f'PDF anexado: {"Sim" if tem_pdf else "Não"}', styles['Normal']))

    doc.build(story)
    return buffer.getvalue()


def gerar_docx_bytes(rad):
    """
    RG-EXP-003: segundo formato de exportação (o primeiro e o PDF).
    Mesma lista de campos de _campos_do_relatorio -- nunca diverge do
    PDF nem da mensagem "Copiar msg".
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    desabilitados = campos_desabilitados()
    campos = _campos_do_relatorio(rad)

    documento = Document()

    titulo = documento.add_heading('RAD — Relatório de Atividade Diária', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitulo = documento.add_paragraph()
    subtitulo.add_run(
        f'RAD N° {rad.numero_rad} — Status: {rad.get_status_display().capitalize()}'
    ).italic = True

    tabela = documento.add_table(rows=0, cols=2)
    tabela.style = 'Light Grid Accent 1'
    tabela.columns[0].width = Pt(140)
    tabela.columns[1].width = Pt(340)

    for chave, rotulo, valor in campos:
        if chave in desabilitados:
            continue
        linha = tabela.add_row()
        celula_rotulo, celula_valor = linha.cells
        celula_rotulo.paragraphs[0].add_run(rotulo).bold = True
        celula_valor.paragraphs[0].add_run(valor)

    mostra_fotos_intervencao = 'fotos_intervencao_verificada' not in desabilitados
    mostra_fotos_acao = 'fotos_acao_realizada' not in desabilitados
    mostra_pdf_anexo = 'pdf' not in desabilitados
    if mostra_fotos_intervencao or mostra_fotos_acao or mostra_pdf_anexo:
        documento.add_heading('Anexos', level=2)
        if mostra_fotos_intervencao:
            qtd = rad.anexos.filter(
                tipo_arquivo='foto', categoria_foto='intervencao_verificada'
            ).count()
            documento.add_paragraph(f'Fotos — Intervenção Verificada: {qtd}')
        if mostra_fotos_acao:
            qtd = rad.anexos.filter(tipo_arquivo='foto', categoria_foto='acao_realizada').count()
            documento.add_paragraph(f'Fotos — Ação Realizada: {qtd}')
        if mostra_pdf_anexo:
            tem_pdf = rad.anexos.filter(tipo_arquivo='pdf').exists()
            documento.add_paragraph(f'PDF anexado: {"Sim" if tem_pdf else "Não"}')

    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()
