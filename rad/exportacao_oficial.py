"""
Exportacao no layout do documento RDA oficial da TRIVIA
(RDA_Relatorio_Trivia_Ajustado_original.docx, fornecido pelo cliente
em 22/07/2026), usado como molde exato (mesmo cabecalho, cores, caixas
de selecao) em vez do layout simples gerado no navegador
(exportar_cliente.js), que so serve para RASCUNHOS locais ainda nao
sincronizados.

O molde tokenizado fica em rad/templates_export/rda_oficial_template.docx
-- e uma copia do documento original com os valores de exemplo trocados
por marcadores {{TOKEN}}. Esta funcao abre o molde, substitui os
marcadores pelos dados reais do RAD, reconstroi a tabela "Atividades
Executadas" com a lista real de Servicos do RAD Digital (o documento
original tinha uma lista fixa de 19 itens especifica de outro
contexto, que nao existe no cadastro atual -- ver conversa de
22/07/2026) e insere as fotos anexadas nos espacos do "Relatorio
Fotografico".

22/07/2026: este e' o UNICO formato de exportacao pos-sincronizacao do
sistema agora (Word). O antigo layout simples (rad/exportacao.py) foi
descontinuado -- ver consulta/views.py.

30/07/2026: o campo Operador CCM deixou de ser um unico texto e passou
a ser dois pares Nome+Hora (Abertura/Entrega). O molde .docx original
so tem UM marcador de texto para isso ({{OPERADOR_CCM}}) -- editar o
proprio arquivo .docx (binario) esta fora do escopo desta mudanca, entao
os dois pares sao combinados num unico texto formatado e inseridos
nesse mesmo marcador. Se no futuro o molde ganhar dois espacos
separados, e so trocar {{OPERADOR_CCM}} por {{OPERADOR_CCM_ABERTURA}}
e {{OPERADOR_CCM_ENTREGA}} no .docx e adicionar os dois tokens
correspondentes no dicionario `substituicoes` abaixo.

Limitacoes conhecidas (a resolver com o cliente):
- "RESPONSAVEL RAD" e preenchido automaticamente com o nome de quem
  sincronizou o RAD (consulta.views.nome_de_quem_preencheu) -- decisao
  do cliente em 22/07/2026, substitui o antigo campo digitado
  manualmente "Equipe TRIVIA".
- "Conclusao e Liberacao da Via" nao tem campo correspondente no RAD
  Digital ainda -- o molde atual tambem nao tem mais essa secao.
- 30/07/2026: o molde ganhou os 4 espacos de foto (antes so 3 eram
  aproveitados) -- ver mapeamento FOTO 1-4 em gerar_docx_oficial_bytes.
- Operador CCM: o RAD Digital guarda 4 informacoes (nome/hora de
  Abertura, nome/hora de Entrega), mas o molde so tem 1 campo, rotulado
  "Op CCM - Abertura". Por isso so a Abertura aparece no Word oficial
  (ver _operador_ccm_abertura_texto) -- a Entrega fica de fora ate o
  molde ganhar um campo proprio para ela.
- Descricoes de foto do VPM001 (Rad.desc_foto_1 a desc_foto_4) e os
  nomes/matriculas dos Colaboradores nao tem espaco reservado no molde
  atual -- os tokens {{NOMES_COLABORADORES}} e
  {{MATRICULAS_COLABORADORES}} existem no codigo mas ficam ociosos ate
  o molde ganhar as secoes correspondentes.
"""
import io
import os

from django.conf import settings
from docx import Document
from docx.shared import Cm
from PIL import Image

CAMINHO_TEMPLATE = os.path.join(
    settings.BASE_DIR, 'rad', 'templates_export', 'rda_oficial_template.docx'
)

# Nomes dos servicos, na mesma ordem/regra usada no formulario
# (rad_form.js::ordenarComOutrosPorUltimo): alfabetica, "Outros" por
# ultimo. Buscados do catalogo real a cada chamada -- se um servico
# novo for cadastrado, aparece aqui automaticamente, sem precisar
# tocar neste arquivo.
ITENS_POR_LINHA_CHECKLIST = 4


def _texto_completo_paragrafo(paragrafo):
    return ''.join(run.text for run in paragrafo.runs)


def _substituir_no_paragrafo(paragrafo, de, para):
    """
    Troca 'de' por 'para' no texto de um paragrafo, mesmo quando o
    texto esta fragmentado em varios <w:r> (comum em docx editados no
    Word). Concatena todos os runs, faz a troca, e realoca o resultado
    no primeiro run -- os demais ficam vazios (nao remove os runs para
    nao perder formatacao caso o paragrafo seja reutilizado).
    """
    texto_atual = _texto_completo_paragrafo(paragrafo)
    if de not in texto_atual:
        return False
    novo_texto = texto_atual.replace(de, para)
    if paragrafo.runs:
        paragrafo.runs[0].text = novo_texto
        for run in paragrafo.runs[1:]:
            run.text = ''
    return True


def _substituir_no_documento(doc, de, para):
    """Aplica _substituir_no_paragrafo em todas as celulas de todas as tabelas."""
    encontrado = False
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas_unicas = []
            ids_vistos = set()
            for celula in linha.cells:
                if id(celula._tc) in ids_vistos:
                    continue
                ids_vistos.add(id(celula._tc))
                celulas_unicas.append(celula)
            for celula in celulas_unicas:
                for paragrafo in celula.paragraphs:
                    if _substituir_no_paragrafo(paragrafo, de, para):
                        encontrado = True
    return encontrado


def _celulas_unicas_da_linha(linha):
    vistas = []
    ids = set()
    for celula in linha.cells:
        if id(celula._tc) in ids:
            continue
        ids.add(id(celula._tc))
        vistas.append(celula)
    return vistas


def _limpar_texto_celula(celula, novo_texto):
    paragrafo = celula.paragraphs[0]
    if paragrafo.runs:
        paragrafo.runs[0].text = novo_texto
        for run in paragrafo.runs[1:]:
            run.text = ''
    else:
        paragrafo.add_run(novo_texto)


def _reconstruir_checklist_servicos(doc, ids_servicos_selecionados):
    """
    Localiza a tabela de "Atividades Executadas" (unica tabela do
    documento) e substitui as linhas de checklist originais (lista
    fixa de 19 itens que nao existe no RAD Digital) pelos SERVICOS
    SELECIONADOS NESSE RAD.

    30/07/2026 (revisado): antes esta funcao listava TODOS os servicos
    do catalogo, marcando "[X]" nos selecionados e "[  ]" nos demais.
    Agora, a pedido do cliente, so os servicos selecionados aparecem
    no documento -- as celulas alem da quantidade selecionada ficam em
    branco, sem lista de "nao selecionados" nenhuma.
    """
    from catalogos.models import CatServico

    servicos_selecionados = list(
        CatServico.objects.filter(id__in=ids_servicos_selecionados).order_by('nome')
    )
    # "Outros" por ultimo, mesmo criterio usado na tela de preenchimento
    sem_outros = [s for s in servicos_selecionados if s.nome != 'Outros']
    outros = [s for s in servicos_selecionados if s.nome == 'Outros']
    servicos_ordenados = sem_outros + outros

    tabela = doc.tables[0]
    linhas_xml = tabela._tbl.findall(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'
    )

    # Acha o indice da primeira linha do checklist procurando pelo
    # texto de uma das celulas (mais resistente a mudanca de posicao
    # do que um indice fixo, caso o molde seja editado no futuro).
    indice_primeira_linha_checklist = None
    indice_ultima_linha_checklist = None
    for indice, linha in enumerate(tabela.rows):
        texto_linha = linha.cells[0].text
        if texto_linha.startswith('[') and indice_primeira_linha_checklist is None:
            indice_primeira_linha_checklist = indice
        elif indice_primeira_linha_checklist is not None and not texto_linha.startswith('['):
            indice_ultima_linha_checklist = indice - 1
            break

    if indice_primeira_linha_checklist is None:
        # Nao achou o padrao esperado -- nao mexe em nada em vez de
        # arriscar corromper o documento.
        return

    linha_modelo_xml = linhas_xml[indice_primeira_linha_checklist]

    # Sempre gera ao menos 1 linha (mesmo com poucos servicos
    # selecionados), pra manter a secao visivel no documento -- as
    # celulas sobrando na linha ficam em branco.
    quantidade_linhas = max(1, -(-len(servicos_ordenados) // ITENS_POR_LINHA_CHECKLIST))  # ceil

    novas_linhas_xml = []
    for indice_linha in range(quantidade_linhas):
        inicio = indice_linha * ITENS_POR_LINHA_CHECKLIST
        grupo = servicos_ordenados[inicio:inicio + ITENS_POR_LINHA_CHECKLIST]
        nova_linha_xml = copy_element(linha_modelo_xml)
        tabela_temp_row = _TrWrapper(nova_linha_xml, tabela)
        celulas = _celulas_unicas_da_linha(tabela_temp_row)
        for posicao, celula in enumerate(celulas):
            if posicao < len(grupo):
                _limpar_texto_celula(celula, grupo[posicao].nome)
            else:
                _limpar_texto_celula(celula, '')
        novas_linhas_xml.append(nova_linha_xml)

    # Remove as linhas antigas do checklist e insere as novas no
    # mesmo lugar. IMPORTANTE (bug corrigido em 30/07/2026): o pai e a
    # linha seguinte precisam ser capturados ANTES do loop de remocao
    # -- linha_de_referencia e uma das linhas removidas, entao chamar
    # .getparent() nela DEPOIS de remover sempre retorna None (o
    # elemento fica orfao). O bug so nao aparecia porque esta funcao
    # nunca tinha sido exercitada de ponta a ponta em producao.
    linha_de_referencia = linhas_xml[indice_primeira_linha_checklist]
    elemento_pai = linha_de_referencia.getparent()
    linha_seguinte = (
        linhas_xml[indice_ultima_linha_checklist + 1]
        if indice_ultima_linha_checklist + 1 < len(linhas_xml)
        else None
    )

    for indice in range(indice_ultima_linha_checklist, indice_primeira_linha_checklist - 1, -1):
        linhas_xml[indice].getparent().remove(linhas_xml[indice])

    if linha_seguinte is not None:
        # Insere logo ANTES do que sobrou depois do checklist (ex.: a
        # linha "3. Equipes Envolvidas"), preservando a posicao visual
        # correta em vez de jogar as linhas novas no fim da tabela.
        for nova_linha in reversed(novas_linhas_xml):
            linha_seguinte.addprevious(nova_linha)
    else:
        # Checklist era a ultima coisa da tabela -- so anexa no fim.
        for nova_linha in novas_linhas_xml:
            elemento_pai.append(nova_linha)


class _TrWrapper:
    """Wrapper minimo para reaproveitar _celulas_unicas_da_linha em um <w:tr> solto (fora da tabela real)."""

    def __init__(self, tr_element, tabela):
        from docx.table import _Row
        self._linha_docx = _Row(tr_element, tabela)

    @property
    def cells(self):
        return self._linha_docx.cells


def copy_element(elemento):
    import copy as copy_module
    return copy_module.deepcopy(elemento)


def _inserir_foto_na_celula(celula, caminho_arquivo, largura_cm=7):
    """
    Substitui o texto de placeholder por uma imagem de verdade,
    redimensionada para caber na largura da coluna. Mantem o
    paragrafo existente (mesma formatacao/alinhamento centralizado do
    molde) em vez de criar um novo.
    """
    paragrafo = celula.paragraphs[0]
    for run in paragrafo.runs:
        run.text = ''
    # Remove paragrafos extras da celula (algumas tem 2, por causa do
    # texto de instrucao em 2 linhas no molde original).
    for p_extra in celula.paragraphs[1:]:
        p_extra._p.getparent().remove(p_extra._p)

    run = paragrafo.add_run()
    run.add_picture(caminho_arquivo, width=Cm(largura_cm))


def _operador_ccm_abertura_texto(rad):
    """
    30/07/2026 (revisado): o campo do molde e literalmente "Op CCM -
    Abertura" (rotulo no Word: "[Nome da Pessoa que abriu]") -- entao
    mostra so o par Abertura (nome + hora), sem misturar com Entrega.
    A Entrega nao tem campo correspondente no molde atual (ver nota no
    docstring do modulo) -- fica de fora do Word oficial por enquanto.
    """
    nome = rad.operador_ccm_abertura_nome
    hora = rad.operador_ccm_abertura_hora
    if not nome:
        return 'N/A'
    return f'{nome} ({hora.strftime("%H:%M")})' if hora else nome


def gerar_docx_oficial_bytes(rad):
    """
    Gera o .docx no layout oficial da TRIVIA para um RAD ja
    sincronizado. Retorna os bytes prontos para download.
    """
    from django.core.files.storage import default_storage

    from consulta.views import nome_de_quem_preencheu

    doc = Document(CAMINHO_TEMPLATE)

    def na(valor):
        return valor if valor not in (None, '') else 'N/A'

    substituicoes = {
        '{{ATIVIDADE}}': na(rad.descricao_tecnica_atividade)[:200],
        '{{OS}}': str(rad.numero_os),
        '{{LOCAL}}': f'{rad.local_inicial.sigla} / {rad.local_final.sigla}',
        '{{RESP_ATIVIDADE}}': na(rad.responsavel_atividade),
        '{{DATA}}': rad.data_preenchimento.strftime('%d/%m/%Y'),
        '{{SA}}': rad.numero_sa,
        '{{SOLICITANTE_SA}}': na(rad.solicitante_sa),
        '{{OPERADOR_CCM}}': _operador_ccm_abertura_texto(rad),
        '{{RESPONSAVEL_RAD}}': nome_de_quem_preencheu(rad),
        '{{H_PROGRAMADO}}': f'{rad.hora_prog_inicio.strftime("%H:%M")} às {rad.hora_prog_termino.strftime("%H:%M")}',
        '{{H_EXECUTADO}}': f'{rad.hora_real_inicio.strftime("%H:%M")} às {rad.hora_real_termino.strftime("%H:%M")}',
        '{{DESCRICAO_TECNICA}}': na(rad.descricao_tecnica_atividade),
        '{{CONCLUSAO}}': (
            'Não há um campo "Conclusão e Liberação da Via" no RAD Digital ainda. '
            'Consulte as Observações Gerais abaixo, se preenchidas.'
        ),
        '{{OBSERVACOES_FOTOS}}': na(rad.observacoes_gerais),
    }

    vias_selecionadas = set(rad.vias.values_list('via__nome', flat=True))
    for numero in (1, 2, 3, 4):
        marcado = f'Via {numero}' in vias_selecionadas
        substituicoes[f'{{{{CB_VIA{numero}}}}}'] = '[X]' if marcado else '[  ]'

    linhas_selecionadas = set(rad.linhas.values_list('linha_id', flat=True))
    for codigo in ('11', '12', '13'):
        marcado = codigo in linhas_selecionadas
        substituicoes[f'{{{{CB_LINHA{codigo}}}}}'] = '[X]' if marcado else '[  ]'

    equipes_selecionadas = set(rad.equipes.values_list('equipe_id', flat=True))
    for codigo, token in (
        ('RA', 'CB_RA'), ('VP', 'CB_VP'), ('CIVIL', 'CB_CIVIL'), ('RESTAB', 'CB_RESTAB'),
        ('MRO', 'CB_MRO'), ('SINAL', 'CB_SINAL'),
    ):
        marcado = codigo in equipes_selecionadas
        substituicoes[f'{{{{{token}}}}}'] = '[X]' if marcado else '[  ]'

    colaboradores = list(rad.colaboradores.all())
    substituicoes['{{NOMES_COLABORADORES}}'] = (
        ' / '.join(c.nome for c in colaboradores) if colaboradores else 'N/A'
    )
    substituicoes['{{MATRICULAS_COLABORADORES}}'] = (
        ' / '.join(c.registro_empresa or 'Participante' for c in colaboradores)
        if colaboradores else 'N/A'
    )

    for token, valor in substituicoes.items():
        _substituir_no_documento(doc, token, valor)

    ids_servicos_selecionados = set(rad.servicos.values_list('servico_id', flat=True))
    _reconstruir_checklist_servicos(doc, ids_servicos_selecionados)

    # Fotos: 30/07/2026 -- o molde ganhou 4 espacos de verdade (antes
    # so tinha 3 aproveitaveis). Mapeamento fixo, casa com o texto
    # "Foto N" que ja vem escrito nas celulas do molde:
    #   FOTO 1 = Intervencao Verificada 1   FOTO 2 = Acao Realizada 1
    #   FOTO 3 = Intervencao Verificada 2   FOTO 4 = Acao Realizada 2
    anexos_intervencao = [a for a in rad.anexos.all() if a.categoria_foto == 'intervencao_verificada']
    anexos_acao = [a for a in rad.anexos.all() if a.categoria_foto == 'acao_realizada']

    tabela = doc.tables[0]
    for linha in tabela.rows:
        celulas = _celulas_unicas_da_linha(linha)
        for celula in celulas:
            if 'FOTO 1' in celula.text and len(anexos_intervencao) > 0:
                with default_storage.open(anexos_intervencao[0].caminho_servidor, 'rb') as arquivo:
                    _inserir_foto_temp(celula, arquivo)
            elif 'FOTO 2' in celula.text and len(anexos_acao) > 0:
                with default_storage.open(anexos_acao[0].caminho_servidor, 'rb') as arquivo:
                    _inserir_foto_temp(celula, arquivo)
            elif 'FOTO 3' in celula.text and len(anexos_intervencao) > 1:
                with default_storage.open(anexos_intervencao[1].caminho_servidor, 'rb') as arquivo:
                    _inserir_foto_temp(celula, arquivo)
            elif 'FOTO 4' in celula.text and len(anexos_acao) > 1:
                with default_storage.open(anexos_acao[1].caminho_servidor, 'rb') as arquivo:
                    _inserir_foto_temp(celula, arquivo)

    buffer_saida = io.BytesIO()
    doc.save(buffer_saida)
    return buffer_saida.getvalue()


def _inserir_foto_temp(celula, arquivo_django):
    """
    python-docx precisa de um caminho de arquivo ou objeto tipo-arquivo
    seekable; grava em um buffer de memoria para nao depender do tipo
    exato de storage configurado (local em dev, outro em producao).
    """
    conteudo = arquivo_django.read()
    buffer_imagem = io.BytesIO(conteudo)
    try:
        imagem = Image.open(buffer_imagem)
        imagem.verify()
    except Exception:
        return  # arquivo corrompido -- nao trava a exportacao inteira por causa de 1 foto
    buffer_imagem.seek(0)
    _inserir_foto_na_celula(celula, buffer_imagem)


class PdfNaoDisponivelError(Exception):
    """
    Levantada quando a exportacao em PDF e chamada mas o servidor nao
    tem o LibreOffice instalado. A view (consulta/views.py) converte
    isso em uma resposta HTTP 503 com mensagem clara -- nunca deve
    virar um erro 500 sem explicacao para quem esta usando o sistema.
    """


def gerar_pdf_oficial_bytes(rad):
    """
    Gera o mesmo layout oficial em PDF, convertendo o .docx (ver
    gerar_docx_oficial_bytes) com o LibreOffice em modo headless.

    22/07/2026: o servidor de producao (Render, plano free, 512MB RAM)
    ainda NAO tem o LibreOffice instalado -- decisao consciente, ver
    conversa com o cliente: cada conversao consome ~150-300MB de RAM,
    o que arrisca derrubar o servico inteiro (OOM) se duas pessoas
    exportarem ao mesmo tempo. Por isso a chamada a esta funcao fica
    atras de um interruptor em Configuracoes
    (CampoFormulario.chave='exportar_pdf_oficial'), desligado por
    padrao -- ver consulta/views.py::exportar_pdf_oficial.

    Esta funcao em si ja esta pronta para quando o LibreOffice for
    instalado (ex.: apos migrar para um plano com mais RAM/disco): so
    precisa adicionar o pacote do LibreOffice ao Dockerfile e ligar o
    interruptor na tela de Configuracoes -- nenhum outro codigo
    precisa mudar.

    Levanta PdfNaoDisponivelError se o binario 'soffice' nao existir
    no PATH do servidor (import tardio de subprocess/shutil para nao
    pesar a inicializacao do modulo quando o PDF nunca e chamado).
    """
    import shutil
    import subprocess
    import tempfile

    if shutil.which('soffice') is None:
        raise PdfNaoDisponivelError(
            'O LibreOffice (comando "soffice") nao esta instalado neste servidor. '
            'A exportacao em PDF nao pode ser gerada ate ele ser instalado.'
        )

    docx_bytes = gerar_docx_oficial_bytes(rad)

    with tempfile.TemporaryDirectory() as pasta_temporaria:
        caminho_docx = os.path.join(pasta_temporaria, f'{rad.numero_rad}.docx')
        with open(caminho_docx, 'wb') as arquivo:
            arquivo.write(docx_bytes)

        resultado = subprocess.run(
            [
                'soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', pasta_temporaria, caminho_docx,
            ],
            capture_output=True,
            timeout=55,  # gunicorn --timeout 60 -- precisa terminar antes disso
        )
        if resultado.returncode != 0:
            raise PdfNaoDisponivelError(
                f'A conversão para PDF falhou: {resultado.stderr.decode(errors="ignore")[:300]}'
            )

        caminho_pdf = os.path.join(pasta_temporaria, f'{rad.numero_rad}.pdf')
        if not os.path.exists(caminho_pdf):
            raise PdfNaoDisponivelError('A conversão para PDF não produziu um arquivo de saída.')

        with open(caminho_pdf, 'rb') as arquivo:
            return arquivo.read()
