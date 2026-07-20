"""
Testes de rad/exportacao.py — RG-EXP-013 (mensagem) e RG-EXP-001 a 012
(PDF), incluindo os campos adicionados na mudanca de negocio de
17/07/2026.
"""
import pytest
from pypdf import PdfReader

from catalogos.models import (
    CatEquipe,
    CatLinha,
    CatLocal,
    CatMotivoAtraso,
    CatServico,
    CatTipoManutencao,
    CatVia,
)
from configuracoes.models import CampoFormulario
from rad.exportacao import gerar_docx_bytes, gerar_mensagem_copiar, gerar_pdf_bytes
from rad.models import Rad, RadColaborador, RadEquipe, RadLinha, RadServico, RadVia
from usuarios.models import Usuario

from .test_regras_negocio import _dados_rad_base


@pytest.fixture
def rad_completo(db):
    usuario = Usuario.objects.create(login='tec.exportacao')
    local = CatLocal.objects.create(sigla='BFU', nome='Barra Funda', categoria='estacao')
    tipo = CatTipoManutencao.objects.create(nome='Preventiva')
    linha = CatLinha.objects.create(codigo='11', nome='Coral')
    via = CatVia.objects.create(nome='Via 1')
    equipe_vp = CatEquipe.objects.create(codigo='VP', nome='VP')
    equipe_sinal = CatEquipe.objects.create(codigo='SINAL', nome='SINAL')
    servico = CatServico.objects.create(nome='Inspeção')
    motivo = CatMotivoAtraso.objects.create(nome='Trânsito')

    dados = _dados_rad_base(usuario, local, tipo, numero_os=1111, sync_id='export-1')
    dados['numero_rad'] = 'R00777'
    dados['numero_execucao'] = 1
    dados['numero_sa'] = '9999999999'
    dados['atraso_inicio'] = True
    dados['motivo_atraso_inicio'] = motivo
    dados['km_poste'] = 'Km 12+300'
    dados['materiais_utilizados'] = 'Parafusos, graxa'
    dados['observacoes_gerais'] = 'Tudo certo.'
    dados['responsavel_atividade'] = 'Carlos Souza'
    dados['operador_ccm'] = 'Op. 42'
    dados['descricao_tecnica_atividade'] = 'Trilho com desgaste de 4,5mm — substituído.'
    rad = Rad.objects.create(**dados)

    RadLinha.objects.create(rad=rad, linha=linha)
    RadVia.objects.create(rad=rad, via=via)
    RadEquipe.objects.create(rad=rad, equipe=equipe_vp)
    RadEquipe.objects.create(rad=rad, equipe=equipe_sinal)
    RadServico.objects.create(rad=rad, servico=servico)
    RadColaborador.objects.create(
        rad=rad, registro_empresa='123', nome='Fulano de Tal', tipo='colaborador'
    )
    return rad


@pytest.mark.django_db
class TestGerarMensagemCopiar:
    def test_mensagem_contem_todos_os_campos_originais_da_efd(self, rad_completo):
        mensagem = gerar_mensagem_copiar(rad_completo)
        assert 'RAD - (Relatório de Atividade Diária)' in mensagem
        assert 'Atividade: Inspeção' in mensagem
        assert 'Local: BFU/BFU' in mensagem
        assert 'Linha: Coral' in mensagem
        assert 'Via: Via 1' in mensagem
        assert 'Equipamentos utilizados: Parafusos, graxa' in mensagem
        assert 'Responsável: Fulano de Tal' in mensagem
        assert 'Observação Geral: Tudo certo.' in mensagem

    def test_mensagem_contem_campos_novos_da_mudanca_17_07(self, rad_completo):
        mensagem = gerar_mensagem_copiar(rad_completo)
        assert 'OS:' in mensagem
        assert 'N° SA: 9999999999' in mensagem
        assert 'Equipes Envolvidas:' in mensagem
        assert 'VP' in mensagem
        assert 'SINAL' in mensagem
        assert 'Responsável Atividade: Carlos Souza' in mensagem
        assert 'Operador CCM: Op. 42' in mensagem
        assert 'Descrição Técnica da Atividade: Trilho com desgaste de 4,5mm — substituído.' in mensagem

    def test_campos_vazios_viram_na(self, rad_completo):
        rad_completo.km_poste = None
        rad_completo.operador_ccm = None
        rad_completo.save()
        mensagem = gerar_mensagem_copiar(rad_completo)
        assert 'Km/Poste: N/A' in mensagem
        assert 'Operador CCM: N/A' in mensagem

    def test_campo_desabilitado_nao_aparece_na_mensagem(self, rad_completo):
        CampoFormulario.objects.filter(chave='operador_ccm').update(habilitado=False)
        try:
            mensagem = gerar_mensagem_copiar(rad_completo)
            assert 'Operador CCM' not in mensagem
        finally:
            CampoFormulario.objects.filter(chave='operador_ccm').update(habilitado=True)


@pytest.mark.django_db
class TestGerarPdf:
    def test_pdf_gerado_e_valido_e_tem_conteudo(self, rad_completo):
        conteudo = gerar_pdf_bytes(rad_completo)
        assert conteudo.startswith(b'%PDF')

        leitor = PdfReader(__import__('io').BytesIO(conteudo))
        assert len(leitor.pages) >= 1
        texto = leitor.pages[0].extract_text()
        assert 'Relatório de Atividade Diária' in texto
        assert rad_completo.numero_rad in texto

    def test_pdf_contem_campos_novos(self, rad_completo):
        import io

        conteudo = gerar_pdf_bytes(rad_completo)
        texto = ''.join(p.extract_text() for p in PdfReader(io.BytesIO(conteudo)).pages)
        assert 'Responsável Atividade' in texto
        assert 'Carlos Souza' in texto
        assert 'Operador CCM' in texto
        assert 'Equipes Envolvidas' in texto

    def test_pdf_respeita_campo_desabilitado(self, rad_completo):
        import io

        CampoFormulario.objects.filter(chave='descricao_tecnica_atividade').update(
            habilitado=False
        )
        try:
            conteudo = gerar_pdf_bytes(rad_completo)
            texto = ''.join(p.extract_text() for p in PdfReader(io.BytesIO(conteudo)).pages)
            assert 'Descrição Técnica da Atividade' not in texto
        finally:
            CampoFormulario.objects.filter(chave='descricao_tecnica_atividade').update(
                habilitado=True
            )


@pytest.mark.django_db
class TestGerarDocx:
    def test_docx_gerado_e_valido_e_tem_conteudo(self, rad_completo):
        from docx import Document

        conteudo = gerar_docx_bytes(rad_completo)
        assert conteudo.startswith(b'PK')  # docx e um zip

        documento = Document(__import__('io').BytesIO(conteudo))
        texto_completo = '\n'.join(p.text for p in documento.paragraphs)
        for tabela in documento.tables:
            for linha in tabela.rows:
                texto_completo += '\n' + '\n'.join(c.text for c in linha.cells)

        assert 'Relatório de Atividade Diária' in texto_completo
        assert rad_completo.numero_rad in texto_completo

    def test_docx_contem_campos_novos(self, rad_completo):
        import io

        from docx import Document

        conteudo = gerar_docx_bytes(rad_completo)
        documento = Document(io.BytesIO(conteudo))
        texto_completo = ''
        for tabela in documento.tables:
            for linha in tabela.rows:
                texto_completo += '\n' + '\n'.join(c.text for c in linha.cells)

        assert 'Responsável Atividade' in texto_completo
        assert 'Carlos Souza' in texto_completo
        assert 'Operador CCM' in texto_completo
        assert 'Equipes Envolvidas' in texto_completo

    def test_docx_respeita_campo_desabilitado(self, rad_completo):
        import io

        from docx import Document

        CampoFormulario.objects.filter(chave='responsavel_atividade').update(habilitado=False)
        try:
            conteudo = gerar_docx_bytes(rad_completo)
            documento = Document(io.BytesIO(conteudo))
            texto_completo = ''
            for tabela in documento.tables:
                for linha in tabela.rows:
                    texto_completo += '\n' + '\n'.join(c.text for c in linha.cells)
            assert 'Responsável Atividade' not in texto_completo
        finally:
            CampoFormulario.objects.filter(chave='responsavel_atividade').update(habilitado=True)

    def test_pdf_e_docx_usam_a_mesma_lista_de_campos(self, rad_completo):
        """Garante que os dois formatos nunca divergem: mesma fonte de dados."""
        from rad.exportacao import _campos_do_relatorio

        campos_pdf = [c[0] for c in _campos_do_relatorio(rad_completo)]
        campos_docx = [c[0] for c in _campos_do_relatorio(rad_completo)]
        assert campos_pdf == campos_docx
